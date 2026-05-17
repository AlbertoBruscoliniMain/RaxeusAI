import subprocess
import sys
import io
import os
from datetime import datetime

# Tutti gli import "pesanti" (chromadb, requests/bs4, pypdf, python-docx,
# ddgs, googlesearch) sono *lazy*: vengono caricati alla prima chiamata della
# tool corrispondente. Questo riduce il cold-start dell'app desktop di ~400ms.
# Il pattern: una funzione `_lazy_*()` che restituisce il modulo o None,
# memoizzando il risultato e l'esito (OK / non installato).

_LAZY_CACHE: dict[str, object] = {}


def _lazy(name: str, importer):
    if name in _LAZY_CACHE:
        return _LAZY_CACHE[name]
    try:
        mod = importer()
    except ImportError:
        mod = None
    _LAZY_CACHE[name] = mod
    return mod


def _lazy_ddgs():
    def _imp():
        from ddgs import DDGS
        return DDGS
    return _lazy("ddgs", _imp)


def _lazy_gsearch():
    def _imp():
        from googlesearch import search as gsearch
        return gsearch
    return _lazy("gsearch", _imp)


def _lazy_requests():
    def _imp():
        import requests
        from bs4 import BeautifulSoup
        return (requests, BeautifulSoup)
    return _lazy("requests", _imp)


def _lazy_pypdf():
    def _imp():
        from pypdf import PdfReader
        return PdfReader
    return _lazy("pypdf", _imp)


def _lazy_docx():
    def _imp():
        from docx import Document as _DocxDocument
        return _DocxDocument
    return _lazy("docx", _imp)


def _lazy_chromadb():
    def _imp():
        import chromadb
        from config import RAG_DB_PATH
        return (chromadb, RAG_DB_PATH)
    return _lazy("chromadb", _imp)

_MAX_READ_CHARS = 8000
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".htm", ".css", ".scss",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env",
    ".csv", ".tsv", ".xml", ".svg",
    ".sh", ".bash", ".zsh", ".ps1", ".bat",
    ".c", ".cpp", ".h", ".hpp", ".java", ".kt", ".go", ".rs", ".rb", ".php",
    ".sql", ".lua", ".swift", ".m", ".r", ".jl",
    ".gitignore", ".dockerfile",
}


def _truncate(text: str) -> str:
    if len(text) > _MAX_READ_CHARS:
        return text[:_MAX_READ_CHARS] + "\n…[file troncato]"
    return text


def _read_docx(path: str) -> str:
    DocxDocument = _lazy_docx()
    if DocxDocument is None:
        return "Errore: libreria python-docx non installata."
    try:
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return _truncate(text) if text.strip() else "(documento vuoto)"
    except Exception as e:
        return f"Errore docx: {e}"


def _read_pdf(path: str) -> str:
    PdfReader = _lazy_pypdf()
    if PdfReader is None:
        return "Errore: libreria pypdf non installata."
    try:
        reader = PdfReader(path)
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        return _truncate(text) if text.strip() else "(nessun testo estraibile)"
    except Exception as e:
        return f"Errore pdf: {e}"


def _read_plain(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read(_MAX_READ_CHARS + 1)
        if not content:
            return "(file vuoto)"
        return _truncate(content)
    except Exception as e:
        return f"Errore: {e}"

def web_search(query: str) -> str:
    DDGS = _lazy_ddgs()
    if DDGS is None:
        return "Errore: libreria ddgs non installata."
    try:
        _old_stderr = sys.stderr
        sys.stderr = io.StringIO()
        try:
            with DDGS() as ddgs:
                results = list(ddgs.text(query, max_results=5))
        finally:
            sys.stderr = _old_stderr
        if not results:
            return "Nessun risultato trovato."
        return "\n\n".join(
            f"**{r['title']}**\n{r['body']}\n{r['href']}" for r in results
        )
    except Exception as e:
        return f"Errore ricerca: {e}"


def read_file(path: str) -> str:
    if not os.path.exists(path):
        return f"File non trovato: {path}"
    if os.path.isdir(path):
        return f"'{path}' è una directory: usa list_dir."
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    return _read_plain(path)


def write_file(path: str, content: str) -> str:
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"File scritto: {path}"
    except Exception as e:
        return f"Errore: {e}"


def run_python(code: str) -> str:
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            capture_output=True, text=True, timeout=10,
            encoding="utf-8", errors="replace",
        )
        output = result.stdout or result.stderr or "(nessun output)"
        return output[:2000]
    except subprocess.TimeoutExpired:
        return "Timeout: codice troppo lento (limite 10s)."
    except Exception as e:
        return f"Errore: {e}"


def _ddgs_urls(query: str, max_results: int = 5) -> list[str]:
    DDGS = _lazy_ddgs()
    if DDGS is None:
        return []
    try:
        _old_stderr = sys.stderr
        sys.stderr = io.StringIO()
        try:
            with DDGS() as ddgs:
                rs = list(ddgs.text(query, max_results=max_results))
        finally:
            sys.stderr = _old_stderr
        return [r["href"] for r in rs if r.get("href")]
    except Exception:
        return []


def google_search(query: str) -> str:
    # Strategia: prima proviamo Google (libreria googlesearch-python). Quando
    # Google blocca il bot (succede spesso, restituisce 0 URL) facciamo
    # fallback ai risultati DuckDuckGo, mantenendo comunque il comportamento
    # "deep fetch" — scarichiamo il contenuto dei primi 2 link in parallelo.
    urls: list[str] = []
    gsearch = _lazy_gsearch()
    if gsearch is not None:
        try:
            _old_stderr = sys.stderr
            sys.stderr = io.StringIO()
            try:
                urls = list(gsearch(query, num_results=5, lang="it"))
            finally:
                sys.stderr = _old_stderr
        except Exception:
            urls = []
    if not urls:
        urls = _ddgs_urls(query, max_results=5)
    if not urls:
        return web_search(query)
    from concurrent.futures import ThreadPoolExecutor
    targets = urls[:2]
    with ThreadPoolExecutor(max_workers=2) as ex:
        fetched = list(ex.map(fetch_url, targets))
    results = [f"[{url}]\n{content}" for url, content in zip(targets, fetched)]
    return "\n\n---\n\n".join(results)


def fetch_url(url: str) -> str:
    pkg = _lazy_requests()
    if pkg is None:
        return "Errore: librerie requests/beautifulsoup4 non installate."
    requests, BeautifulSoup = pkg
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = requests.get(url, headers=headers, timeout=8)
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = " ".join(soup.get_text(separator=" ").split())
        return text[:3000]
    except Exception as e:
        return f"Errore nel fetch della pagina: {e}"


def get_datetime() -> str:
    return datetime.now().strftime("%A %d %B %Y, %H:%M:%S")


def list_dir(path: str = ".") -> str:
    try:
        entries = sorted(os.listdir(path))
        if not entries:
            return "(directory vuota)"
        return "\n".join(entries)
    except FileNotFoundError:
        return f"Directory non trovata: {path}"
    except Exception as e:
        return f"Errore: {e}"


def append_file(path: str, content: str) -> str:
    try:
        with open(path, "a", encoding="utf-8") as f:
            f.write(content)
        return f"Contenuto aggiunto a: {path}"
    except Exception as e:
        return f"Errore: {e}"


def read_pdf(path: str) -> str:
    if not os.path.exists(path):
        return f"File non trovato: {path}"
    return _read_pdf(path)


def rag_search(query: str, n_results: int = 4) -> str:
    pkg = _lazy_chromadb()
    if pkg is None:
        return "Errore: chromadb non installato. Esegui: pip install chromadb"
    chromadb, RAG_DB_PATH = pkg
    try:
        client = chromadb.PersistentClient(path=RAG_DB_PATH)
        try:
            collection = client.get_collection("documents")
        except Exception:
            return "Nessun documento indicizzato. Usa rag_index.py per indicizzare dei file."
        results = collection.query(query_texts=[query], n_results=n_results)
        docs = results["documents"][0]
        metas = results["metadatas"][0]
        if not docs:
            return "Nessun risultato rilevante trovato nei documenti indicizzati."
        parts = []
        for doc, meta in zip(docs, metas):
            source = meta.get("source", "sconosciuto")
            parts.append(f"[Fonte: {source}]\n{doc}")
        return "\n\n---\n\n".join(parts)
    except Exception as e:
        return f"Errore RAG: {e}"


def wikipedia_search(query: str, lang: str = "it") -> str:
    pkg = _lazy_requests()
    if pkg is None:
        return "Errore: libreria requests non installata."
    requests, _ = pkg
    # Wikipedia REST richiede uno User-Agent identificativo (policy 2024+),
    # altrimenti risponde 403.
    headers = {
        "User-Agent": "RaxeusAI/1.0 (https://github.com/AlbertoBruscoliniMain/RaxeusAI)",
        "Accept": "application/json",
    }
    try:
        url = f"https://{lang}.wikipedia.org/api/rest_v1/page/summary/{query.replace(' ', '_')}"
        resp = requests.get(url, headers=headers, timeout=8)
        if resp.status_code == 404:
            return f"Nessuna pagina Wikipedia trovata per: {query}"
        if resp.status_code != 200:
            return f"Errore Wikipedia: HTTP {resp.status_code}"
        data = resp.json()
        title = data.get("title", "")
        extract = data.get("extract", "Nessun contenuto disponibile.")
        return f"**{title}**\n\n{extract}"
    except Exception as e:
        return f"Errore Wikipedia: {e}"


TOOLS = {
    "google_search": google_search,
    "web_search": web_search,
    "fetch_url": fetch_url,
    "read_file": read_file,
    "write_file": write_file,
    "append_file": append_file,
    "list_dir": list_dir,
    "read_pdf": read_pdf,
    "wikipedia_search": wikipedia_search,
    "run_python": run_python,
    "get_datetime": get_datetime,
    "rag_search": rag_search,
}

TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "google_search",
            "description": "Cerca su Google. È il motore di ricerca principale — usalo sempre come prima scelta quando hai bisogno di cercare informazioni su internet.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "La query di ricerca"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Legge il contenuto testuale di una pagina web dato il suo URL. Usalo dopo google_search per leggere il contenuto delle pagine trovate e ottenere informazioni aggiornate.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL della pagina da leggere"}
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Cerca informazioni su internet tramite DuckDuckGo. Usalo solo se google_search fallisce o non dà risultati.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "La query di ricerca"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Legge qualsiasi file dal filesystem: testo semplice (txt, md, codice sorgente, json, yaml, csv, log, ecc.), documenti Word (.docx) o PDF. Riconosce automaticamente il formato dall'estensione. L'output è troncato a 8000 caratteri.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Percorso del file"}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Scrive o sovrascrive un file sul filesystem.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Percorso del file"},
                    "content": {"type": "string", "description": "Contenuto da scrivere"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_python",
            "description": "Esegue codice Python e restituisce l'output.",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "Codice Python da eseguire"}
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_datetime",
            "description": "Restituisce la data e l'ora corrente.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": "Elenca i file e le cartelle in una directory. Usalo per esplorare il filesystem prima di leggere o scrivere file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Percorso della directory (default: directory corrente)"}
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "append_file",
            "description": "Aggiunge testo in fondo a un file esistente senza sovrascriverlo. Utile per log, diari, note progressive.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Percorso del file"},
                    "content": {"type": "string", "description": "Testo da aggiungere in fondo al file"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Legge e restituisce il testo estratto da un file PDF. Usalo quando l'utente chiede di analizzare o riassumere un PDF.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Percorso del file PDF"}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "rag_search",
            "description": "Cerca nei documenti personali indicizzati (note, appunti, PDF, file di testo). Usalo quando l'utente chiede qualcosa che potrebbe trovarsi nei suoi file personali salvati localmente.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Cosa cercare nei documenti"},
                    "n_results": {"type": "integer", "description": "Numero di risultati da restituire (default: 4)"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "wikipedia_search",
            "description": "Cerca su Wikipedia e restituisce il sommario della pagina. Usalo per definizioni, biografie, concetti storici o scientifici.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Titolo o argomento da cercare"},
                    "lang": {"type": "string", "description": "Codice lingua Wikipedia (default: 'it' per italiano, 'en' per inglese)"},
                },
                "required": ["query"],
            },
        },
    },
]


def execute_tool(name: str, args: dict) -> str:
    if name not in TOOLS:
        return f"Tool '{name}' non trovato."
    try:
        return str(TOOLS[name](**args))
    except Exception as e:
        return f"Errore nell'esecuzione di '{name}': {e}"
